import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

HEADER_COLOR = RGBColor(0x7F, 0x3D, 0x1E)

doc = docx.Document('APM_Concept_Notes_Manual.docx')
new_doc = docx.Document()

# ---------------------------------------------------------
# 1. FRONT COVER PAGE (IBM SkillsBuild 2026 & Team LOGIC LEGENDS)
# ---------------------------------------------------------
p_inst = new_doc.add_paragraph()
p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_inst = p_inst.add_run("IBM SKILLSBUILD ACADEMIC INTERNSHIPS 2026\nAICTE | IBM SkillsBuild AI Automation & Intelligent Solutions Internship | BharatCares\n")
run_inst.bold = True
run_inst.font.size = Pt(16)
run_inst.font.color.rgb = HEADER_COLOR

p_title = new_doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("PROJECT DOCUMENTATION & CONCEPT NOTE\nAI-Driven Non-Invasive Diabetes Likelihood Prediction System\n")
run_title.bold = True
run_title.font.size = Pt(16)
run_title.font.color.rgb = HEADER_COLOR

p_meta = new_doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_meta = p_meta.add_run(
    "Team Name: LOGIC LEGENDS  |  Unique Team ID: IBMBH00174\n"
    "College: SHRI SHANKARACHARYA TECHNICAL CAMPUS BHILAI (Pincode: 490020)\n"
    "Track: Artificial Intelligence & Machine Learning (SDG Goal 3)\n"
)
run_meta.font.size = Pt(14)

table = new_doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
headers = ["S.No.", "Team Member Name", "Designation / Role", "Registered Email ID"]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(14)

members = [
    ("1", "Yukesh Choudhary", "Team Leader", "yukeshme@gmail.com"),
    ("2", "Tushar Kumar", "Co-Lead", "tusharsahu757@gmail.com"),
    ("3", "Kuldeepak Sahu", "Co-Lead", "kuldeepaksahu8717@gmail.com"),
    ("4", "Kamlesh Biswas", "Co-Lead", "kamleshbiswas96@gmail.com"),
    ("5", "Arpan Raj Kumar", "Team Member", "arpanrajkumar59@gmail.com"),
    ("6", "Lekhram Tarak", "Team Member", "lekhramtarak30@gmail.com"),
]

for sno, name, role, email in members:
    row_cells = table.add_row().cells
    row_cells[0].text = sno
    row_cells[1].text = name
    row_cells[2].text = role
    row_cells[3].text = email
    for cell in row_cells:
        cell.paragraphs[0].runs[0].font.size = Pt(14)

p_pb = new_doc.add_paragraph()
p_pb.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

# ---------------------------------------------------------
# 2. COPY ORIGINAL PARAGRAPHS & APPLY EXPANSIONS & STYLING
# ---------------------------------------------------------
heading_pattern = re.compile(
    r'^(Section \d+:|CHAPTER \d+:|Chapter \d+:|PART \d+:|Part \d+:|PART \d+|CHAPTER \d+|Section \d+|\d+\.\d+\s+[A-Z]|BIBLIOGRAPHY|IBM SKILLSBUILD|PROJECT DOCUMENTATION)',
    re.IGNORECASE
)

for element in doc.element.body:
    new_doc.element.body.append(element)

# Apply global font sizing (16pt centered headings, 14pt content)
for p in new_doc.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    if heading_pattern.match(text) or text.startswith("==========") or text.startswith("CHAPTER ") or text.startswith("Section "):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = HEADER_COLOR
    else:
        for run in p.runs:
            run.font.size = Pt(14)

for t in new_doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if p.text.strip():
                    for run in p.runs:
                        run.font.size = Pt(14)

new_doc.save('APM_Concept_Notes_Manual.docx')
print("SUCCESS: Fully rebuilt APM_Concept_Notes_Manual.docx with cover page, headings=16pt, content=14pt!")
